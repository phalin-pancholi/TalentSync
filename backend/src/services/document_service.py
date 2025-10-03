"""
Document service for TalentSync backend
"""
import os
import io
from datetime import datetime
from typing import Optional
from bson import ObjectId
from fastapi import HTTPException, UploadFile
import PyPDF2
from docx import Document as DocxDocument

from ..models.document import Document, DocumentResponse
from ..services.db_service import get_database


class DocumentService:
    def __init__(self):
        self._db = None
        self._collection = None
    
    @property
    def db(self):
        if self._db is None:
            self._db = get_database()
        return self._db
    
    @property
    def collection(self):
        if self._collection is None:
            self._collection = self.db.documents
        return self._collection
        
    async def create_document(self, candidate_id: str, file: UploadFile) -> str:
        """Create a new document from uploaded file"""
        try:
            # Validate file type
            file_extension = file.filename.split('.')[-1].upper() if file.filename else ""
            if file_extension not in ['PDF', 'DOCX', 'TXT']:
                raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF, DOCX, and TXT are allowed.")
            
            # Extract text content
            content = await file.read()
            text_content = await self._extract_text(content, file_extension)
            
            # Save file to disk
            file_path = await self._save_file(content, file.filename, candidate_id)
            
            # Create document record
            document_data = {
                'candidate_id': ObjectId(candidate_id),
                'file_name': file.filename,
                'file_type': file_extension,
                'content_text': text_content,
                'raw_file_path': file_path,
                'upload_date': datetime.now(timezone.utc)
            }
            
            result = await self.collection.insert_one(document_data)
            return str(result.inserted_id)
            
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error creating document: {str(e)}")

    async def get_document(self, document_id: str) -> Optional[Document]:
        """Get a document by ID"""
        try:
            document = await self.collection.find_one({"_id": ObjectId(document_id)})
            if document:
                return Document(**document)
            return None
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error retrieving document: {str(e)}")

    async def get_document_by_candidate(self, candidate_id: str) -> Optional[Document]:
        """Get a document by candidate ID"""
        try:
            document = await self.collection.find_one({"candidate_id": ObjectId(candidate_id)})
            if document:
                return Document(**document)
            return None
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error retrieving document by candidate: {str(e)}")

    async def delete_document(self, document_id: str) -> bool:
        """Delete a document"""
        try:
            # Get document to find file path
            document = await self.get_document(document_id)
            if document:
                # Delete file from disk
                try:
                    if os.path.exists(document.raw_file_path):
                        os.remove(document.raw_file_path)
                except OSError:
                    pass  # Continue even if file deletion fails
                
                # Delete from database
                result = await self.collection.delete_one({"_id": ObjectId(document_id)})
                return result.deleted_count > 0
            return False
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

    async def _extract_text(self, content: bytes, file_type: str) -> str:
        """Extract text content from file based on type"""
        try:
            if file_type == 'TXT':
                return content.decode('utf-8')
            
            elif file_type == 'PDF':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            
            elif file_type == 'DOCX':
                doc = DocxDocument(io.BytesIO(content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip()
            
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            # If text extraction fails, return empty string rather than failing
            print(f"Warning: Text extraction failed for {file_type}: {str(e)}")
            return ""

    async def _save_file(self, content: bytes, filename: str, candidate_id: str) -> str:
        """Save file to disk and return file path"""
        try:
            # Create candidate-specific directory
            candidate_dir = os.path.join(self.upload_directory, candidate_id)
            os.makedirs(candidate_dir, exist_ok=True)
            
            # Generate unique filename with timestamp
            timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            file_extension = filename.split('.')[-1] if filename and '.' in filename else 'bin'
            unique_filename = f"{timestamp}_{filename}"
            
            file_path = os.path.join(candidate_dir, unique_filename)
            
            # Write file
            with open(file_path, 'wb') as f:
                f.write(content)
            
            return file_path
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")

    def to_response(self, document: Document) -> DocumentResponse:
        """Convert Document to DocumentResponse"""
        return DocumentResponse(
            id=str(document.id),
            candidate_id=str(document.candidate_id),
            file_name=document.file_name,
            file_type=document.file_type,
            content_text=document.content_text,
            raw_file_path=document.raw_file_path,
            upload_date=document.upload_date
        )

    def generate_profile_summary_pdf(self, candidate_name: str, profile_summary: str) -> bytes:
        """
        Generate a PDF from profile summary text with proper formatting
        Converts markdown to properly formatted PDF with headings, sections, and bullets
        """
        try:
            import re
            
            # Parse and format markdown content for PDF
            def parse_markdown_to_pdf_elements(text):
                if not text:
                    return [{"type": "text", "content": "No summary available", "font_size": 10}]
                
                elements = []
                lines = text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        elements.append({"type": "spacing", "content": "", "font_size": 10})
                        continue
                    
                    # Headers (# ## ### etc.)
                    if line.startswith('#'):
                        level = len(line) - len(line.lstrip('#'))
                        content = line.lstrip('# ').strip()
                        if level == 1:
                            elements.append({"type": "heading1", "content": content, "font_size": 14})
                        elif level == 2:
                            elements.append({"type": "heading2", "content": content, "font_size": 12})
                        else:
                            elements.append({"type": "heading3", "content": content, "font_size": 11})
                    
                    # Bullet points (- or * or •)
                    elif line.startswith(('- ', '* ', '• ')):
                        content = line[2:].strip()
                        # Remove markdown formatting from bullet content
                        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Bold
                        content = re.sub(r'\*(.*?)\*', r'\1', content)      # Italic
                        elements.append({"type": "bullet", "content": f"* {content}", "font_size": 10})
                    
                    # Numbered lists (1. 2. etc.)
                    elif re.match(r'^\d+\.\s+', line):
                        content = re.sub(r'^\d+\.\s+', '', line)
                        # Remove markdown formatting
                        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
                        content = re.sub(r'\*(.*?)\*', r'\1', content)
                        number = re.match(r'^(\d+)\.', line).group(1)
                        elements.append({"type": "numbered", "content": f"{number}. {content}", "font_size": 10})
                    
                    # Regular text (remove markdown formatting)
                    else:
                        # Remove bold, italic, and other markdown
                        content = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
                        content = re.sub(r'\*(.*?)\*', r'\1', content)   # Italic
                        content = re.sub(r'`([^`]+)`', r'\1', content)   # Code
                        content = re.sub(r'__(.*?)__', r'\1', content)   # Bold alternative
                        content = re.sub(r'_(.*?)_', r'\1', content)     # Italic alternative
                        
                        if content.strip():
                            elements.append({"type": "text", "content": content, "font_size": 10})
                
                return elements
            
            # Escape special characters for PDF
            def escape_pdf_string(text):
                if not text:
                    return ""
                # Replace problematic Unicode characters with ASCII equivalents
                text = text.replace('•', '* ')  # Replace bullet with asterisk
                text = text.replace('─', '-')   # Replace em-dash with hyphen
                text = text.replace('–', '-')   # Replace en-dash with hyphen
                text = text.replace('—', '-')   # Replace em-dash with hyphen
                text = text.replace('"', '"')   # Replace smart quotes
                text = text.replace('"', '"')
                text = text.replace(''', "'")
                text = text.replace(''', "'")
                # Basic PDF escaping
                return text.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')
            
            # Split long text into lines with proper wrapping
            def wrap_text_for_pdf(text, max_chars=65):
                if not text:
                    return [""]
                
                words = text.split()
                lines = []
                current_line = ""
                
                for word in words:
                    if len(current_line + " " + word) <= max_chars:
                        current_line += " " + word if current_line else word
                    else:
                        if current_line:
                            lines.append(current_line)
                        current_line = word
                
                if current_line:
                    lines.append(current_line)
                
                return lines if lines else [""]
            
            # Parse the profile summary
            elements = parse_markdown_to_pdf_elements(profile_summary)
            
            # Create PDF content stream
            page_content = "BT\n"
            page_content += "50 750 Td\n"  # Start position
            
            # Title (Bold)
            page_content += "/F2 18 Tf\n"  # Use F2 for bold font
            page_content += f"({escape_pdf_string('PROFILE SUMMARY')}) Tj\n"
            page_content += "0 -25 Td\n"
            
            # Candidate name (Bold)
            page_content += "/F2 14 Tf\n"  # Use F2 for bold font
            page_content += f"({escape_pdf_string(candidate_name or 'Unknown Candidate')}) Tj\n"
            page_content += "0 -20 Td\n"
            
            # Add separator line
            page_content += "/F1 10 Tf\n"
            page_content += f"({escape_pdf_string('-' * 60)}) Tj\n"
            page_content += "0 -15 Td\n"
            
            y_position = 680  # Track vertical position
            
            # Process each element
            for element in elements:
                if y_position < 50:  # Near bottom of page
                    break
                
                element_type = element["type"]
                content = element["content"]
                font_size = element["font_size"]
                
                if element_type == "spacing":
                    page_content += "0 -8 Td\n"
                    y_position -= 8
                elif element_type == "heading1":
                    page_content += "/F2 14 Tf\n"  # Bold font for headings
                    page_content += f"({escape_pdf_string(content)}) Tj\n"
                    page_content += "0 -18 Td\n"
                    y_position -= 18
                elif element_type == "heading2":
                    page_content += "/F2 12 Tf\n"  # Bold font for headings
                    page_content += f"({escape_pdf_string(content)}) Tj\n"
                    page_content += "0 -16 Td\n"
                    y_position -= 16
                elif element_type == "heading3":
                    page_content += "/F2 11 Tf\n"  # Bold font for headings
                    page_content += f"({escape_pdf_string(content)}) Tj\n"
                    page_content += "0 -14 Td\n"
                    y_position -= 14
                elif element_type in ["bullet", "numbered"]:
                    # Move right for bullet/numbered list indentation
                    page_content += "15 0 Td\n"  # Indent bullet points from left margin
                    page_content += "/F1 10 Tf\n"
                    # Wrap bullet content if too long (reduce width due to indentation)
                    wrapped_lines = wrap_text_for_pdf(content, 55)
                    for i, line in enumerate(wrapped_lines):
                        if i == 0:
                            # First line of bullet/numbered item
                            page_content += f"({escape_pdf_string(line)}) Tj\n"
                        else:
                            # Continuation lines - additional indent for wrapped text
                            page_content += "15 0 Td\n"  # Additional indent for continuation
                            page_content += f"({escape_pdf_string(line)}) Tj\n"
                            page_content += "-15 0 Td\n"  # Move back to bullet indent level
                        page_content += "0 -12 Td\n"
                        y_position -= 12
                        if y_position < 50:
                            break
                    # Move back to left margin after bullet point
                    page_content += "-15 0 Td\n"
                else:  # Regular text
                    page_content += "/F1 10 Tf\n"
                    # Wrap regular text
                    wrapped_lines = wrap_text_for_pdf(content, 65)
                    for line in wrapped_lines:
                        page_content += f"({escape_pdf_string(line)}) Tj\n"
                        page_content += "0 -12 Td\n"
                        y_position -= 12
                        if y_position < 50:
                            break
            
            # Footer
            if y_position >= 80:
                page_content += "0 -10 Td\n"
                page_content += "/F1 8 Tf\n"
                footer_text = f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')} UTC by TalentSync"
                page_content += f"({escape_pdf_string(footer_text)}) Tj\n"
            
            page_content += "ET"
            
            # Basic PDF structure with both regular and bold fonts
            pdf_content = f"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj

2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj

3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
/F2 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica-Bold
>>
>>
>>
/Contents 4 0 R
>>
endobj

4 0 obj
<<
/Length {len(page_content)}
>>
stream
{page_content}
endstream
endobj

xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000356 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
{456 + len(page_content)}
%%EOF"""
            
            return pdf_content.encode('utf-8')
            
        except Exception as e:
            # Fallback: return a very basic PDF with error message
            def simple_escape(text):
                if not text:
                    return ""
                return text.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')
            
            error_msg = f"Error generating PDF: {str(e)}"
            candidate_safe = simple_escape(candidate_name or "Unknown")[:50]  # Limit length
            error_safe = simple_escape(error_msg)[:100]  # Limit length
            
            basic_pdf = f"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Resources<</Font<</F1<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>/F2<</Type/Font/Subtype/Type1/BaseFont/Helvetica-Bold>>>>>>/Contents 4 0 R>>endobj
4 0 obj<</Length 120>>stream
BT
/F2 12 Tf
50 750 Td
(Profile Summary - {candidate_safe}) Tj
0 -20 Td
/F1 10 Tf
({error_safe}) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000356 00000 n 
trailer<</Size 5/Root 1 0 R>>
startxref
440
%%EOF"""
            return basic_pdf.encode('utf-8')